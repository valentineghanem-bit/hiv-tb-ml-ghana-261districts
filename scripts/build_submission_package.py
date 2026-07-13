"""Build the Epidemiology & Infection submission package."""
from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "submission_package"
MANUSCRIPT_SRC = ROOT / "manuscript" / "HIV_TB_Ghana_261Districts_Manuscript_DRAFT.md"
MANUSCRIPT_MD = PKG / "Ghanem_EI_Manuscript_Submission.md"
MANUSCRIPT_DOCX = PKG / "Ghanem_EI_Manuscript_Submission.docx"


def read_manuscript() -> str:
    text = MANUSCRIPT_SRC.read_text(encoding="utf-8")
    text = text.replace("¹", "1")
    text = text.replace("## Acknowledgements\n\nNone.", "## Acknowledgements\n\nThe author thanks the DHS Programme, WHO and Ghana Statistical Service for access to aggregate public-health data sources.")
    if "## Use of AI Tools" not in text:
        text = text.replace(
            "## Author Contributions\n\nV.G.G.:",
            (
                "## Use of AI Tools\n\n"
                "I used OpenAI Codex to help check consistency across the manuscript, data tables, dashboard and poster, and to help format the submission files. I reviewed and corrected the outputs and take responsibility for the analysis, interpretation and final wording.\n\n"
                "## Author Contributions\n\nV.G.G.:"
            ),
        )
    text = text.replace("## References (Vancouver, numbered by first appearance in Results/Methods)", "## References")
    return text


def add_table_from_frame(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_paragraph(title, style="Caption")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, float):
                cells[i].text = f"{value:.3f}".rstrip("0").rstrip(".")
            else:
                cells[i].text = str(value)
    format_table(table)
    doc.add_paragraph()


def format_table(table) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True


def write_simple_docx(path: Path, text: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)
    doc.save(path)


def write_table_docx(path: Path, title: str, intro: str, df: pd.DataFrame, footer: str = "") -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(4)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(intro)
    add_table_from_frame(doc, df, "")
    if footer:
        note = doc.add_paragraph(footer)
        note.runs[0].italic = True
    doc.save(path)


def add_results_tables(doc: Document) -> None:
    moran = pd.read_csv(ROOT / "outputs" / "tables" / "global_morans_I.csv")
    add_table_from_frame(doc, moran, "Table 1. Global Moran's I for district-level HIV, TB and system indicators.")

    gwr_global = pd.read_csv(ROOT / "outputs" / "tables" / "gwr_global_fit.csv")
    gwr_summary = pd.read_csv(ROOT / "outputs" / "tables" / "gwr_summary.csv")
    add_table_from_frame(doc, gwr_global, "Table 2a. GWR global fit diagnostics.")
    add_table_from_frame(doc, gwr_summary, "Table 2b. GWR coefficient summary.")

    ml10 = pd.read_csv(ROOT / "outputs" / "tables" / "ml_10fold_cv_results.csv")
    mlsp = pd.read_csv(ROOT / "outputs" / "tables" / "ml_spatial_cv_results.csv")
    ml = ml10.merge(mlsp, on="Model", how="left")
    add_table_from_frame(doc, ml, "Table 3. Machine-learning performance under random 10-fold and spatial cross-validation.")

    provenance = pd.DataFrame(
        [
            ["Ghana Statistical Service 2021 Census", "District socioeconomic and demographic indicators", "District", "2021"],
            ["WHO Global Health Observatory / Global TB Programme", "TB incidence, TB-HIV, ART and treatment indicators", "National time series mapped into ecological frame", "2013-2024"],
            ["Ghana DHS", "HIV prevalence and behavioural/testing/attitudinal indicators", "Regional", "2003 round for behavioural inputs used in pipeline"],
            ["Ghana district GeoJSON", "260 unique district polygons plus Guan shared-polygon convention", "Spatial geometry", "2021/2018 administrative frame"],
        ],
        columns=["Source", "Variables used", "Scale", "Date/round"],
    )
    add_table_from_frame(doc, provenance, "Table 4. Data provenance and scale disclosure.")


def add_figures(doc: Document) -> None:
    figures = [
        ("Figure 1. Disease burden choropleths.", "Figure_1_disease_burden.png"),
        ("Figure 2. Global Moran's I scatterplots.", "Figure_3_morans_I.png"),
        ("Figure 3. LISA, bivariate LISA and Getis-Ord Gi* clusters.", "Figure_2_spatial_clusters.png"),
        ("Figure 4. GWR local coefficient maps.", "Figure_8_gwr_coefficients.png"),
        ("Figure 5. ML model comparison: random 10-fold versus spatial CV.", "Figure_4_ml_performance.png"),
        ("Figure 6. SHAP feature importance.", "Figure_5_shap_importance.png"),
        ("Figure 6b. SHAP beeswarm.", "Figure_5b_shap_beeswarm.png"),
        ("Figure 7. Stacked-ensemble risk map.", "Figure_6_ml_risk_map.png"),
        ("Figure 8. Predictor correlation matrix.", "Figure_7_correlation.png"),
    ]
    for caption, file_name in figures:
        doc.add_paragraph(caption, style="Caption")
        path = ROOT / "outputs" / "figures" / file_name
        if path.exists():
            doc.add_picture(str(path), width=Inches(6.4))
        doc.add_paragraph()


def md_to_docx(text: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    styles["Normal"].paragraph_format.space_after = Pt(0)

    in_table_placeholder = False
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line == "---":
            continue
        if line.startswith("## Tables"):
            doc.add_heading("Tables", level=1)
            add_results_tables(doc)
            in_table_placeholder = True
            continue
        if line.startswith("## Figures"):
            doc.add_heading("Figures", level=1)
            add_figures(doc)
            in_table_placeholder = False
            continue
        if in_table_placeholder:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line, style="List Number")
        else:
            clean = line.replace("**", "").replace("*", "")
            doc.add_paragraph(clean)

    doc.save(MANUSCRIPT_DOCX)


def write_cover_letter() -> None:
    text = """Valentine Golden Ghanem
Ghana COCOBOD Cocoa Clinic, Accra, Ghana
valentineghanem@gmail.com | ORCID: 0009-0002-8332-0220

13 July 2026

Dear Editors,

I am submitting the manuscript "District-Level Spatial and Machine-Learning Analysis of HIV-TB Co-infection in Ghana" as an Original Paper for Epidemiology & Infection.

The paper is a district-level infectious-disease epidemiology study with direct public-health relevance. Its main point is deliberately modest. HIV-TB co-infection clusters spatially across Ghana, and ordinary random-fold validation overstates how well a model will travel to unseen geography: LightGBM AUC falls from 0.998 under 10-fold cross-validation to 0.798 under leave-one-region-out spatial cross-validation. That gap matters for surveillance studies built on geographically autocorrelated health data.

The manuscript uses aggregate district, regional and national indicators from Ghana Statistical Service, WHO/GHO and DHS-derived sources. The analysis combines global and local Moran's I, bivariate LISA, Getis-Ord Gi*, geographically weighted regression, Random Forest, XGBoost, LightGBM and SHAP interpretation. The limitations around DHS behavioural variables and the Guan shared-polygon convention are stated directly, rather than left for reviewers to infer.

This manuscript is original, is not under consideration elsewhere and has not been published previously. The author has read and approved the submitted version. There are no conflicts of interest and no specific funding to declare. Code, master data and reproducible outputs are available at https://github.com/valentineghanem-bit/hiv-tb-ml-ghana-260districts.

I used OpenAI Codex to help check consistency across the manuscript, data tables, dashboard and poster, and to help format the submission files. I reviewed and corrected the outputs and take responsibility for all submitted content.

Sincerely,

Valentine Golden Ghanem
"""
    write_simple_docx(PKG / "Ghanem_EI_Cover_Letter.docx", text)


def write_checklists() -> None:
    checklist = """# Epidemiology & Infection Submission Checklist

- Article type: Original Paper.
- Summary: 190 words, within the 150-200 word journal range.
- Main text: approximately 3,530 words, within the 2,000-4,000 word journal range.
- References: 20, below the journal's recommended maximum of 40.
- Reporting: STROBE declared and checklist supplied.
- Data availability: public code/data repository supplied; raw DHS access limitation disclosed.
- Ethics: secondary aggregate ecological analysis; no individual-level patient data accessed.
- AI declaration: included in manuscript and cover letter.
- Visual companions: dashboard and poster regenerated through Bespoke HI-EI generator.
- Sensitivity checks: structural/TB/system-only spatial CV and Guan/Krachi East shared-polygon check supplied in `outputs/tables/submission_sensitivity_checks.csv`.
- Repository: aligned to https://github.com/valentineghanem-bit/hiv-tb-ml-ghana-260districts.
"""
    write_simple_docx(PKG / "EI_Submission_Checklist.docx", checklist)

    strobe_rows = [
        ["1a", "Title identifies the design or analytic frame", "Title page", "Met"],
        ["1b", "Balanced summary of methods and findings", "Summary", "Met"],
        ["2", "Scientific background and rationale", "Introduction, paragraphs 1-4", "Met"],
        ["3", "Specific objectives", "End of Introduction", "Met"],
        ["4", "Study design stated early", "Methods 2.1", "Met"],
        ["5", "Setting, locations and study period", "Methods 2.1-2.2; Table 4", "Met"],
        ["6", "Eligibility/unit of analysis", "Methods 2.1; Results 3.1", "Met"],
        ["7", "Variables and outcomes defined", "Methods 2.3", "Met"],
        ["8", "Data sources and measurement", "Methods 2.2; Table 4", "Met"],
        ["9", "Bias and design limitations described", "Discussion 4.5", "Met"],
        ["10", "Study size explained", "Results 3.1", "Met"],
        ["11", "Quantitative variables handled transparently", "Methods 2.3-2.6", "Met"],
        ["12a", "Statistical methods", "Methods 2.4-2.6", "Met"],
        ["12b", "Subgroup/spatial methods", "Methods 2.4-2.5", "Met"],
        ["12c", "Missing data approach", "Methods 2.3; repository pipeline", "Met"],
        ["12d", "Sensitivity analyses", "Methods 2.7; Supplementary sensitivity table", "Added"],
        ["13", "Participant flow/unit inclusion", "Results 3.1", "Met"],
        ["14", "Descriptive data", "Results 3.1-3.3; Tables 1-4", "Met"],
        ["15", "Outcome data", "Results 3.2-3.7", "Met"],
        ["16", "Main results with precision where available", "Results 3.2-3.7; Tables 1-3", "Met"],
        ["17", "Other analyses and robustness checks", "Results 3.8; sensitivity table", "Added"],
        ["18", "Key results interpreted against objectives", "Discussion 4.1-4.4", "Met"],
        ["19", "Limitations", "Discussion 4.5", "Met"],
        ["20", "Generalisability", "Discussion 4.2, 4.5 and 4.6", "Met"],
        ["21", "Overall interpretation", "Discussion 4.6", "Met"],
        ["22", "Funding", "Financial Support", "Met"],
    ]
    write_table_docx(
        PKG / "STROBE_Location_Checklist.docx",
        "STROBE Location Checklist",
        "Location map for the Epidemiology & Infection submission. This is a location checklist, not a substitute for the full STROBE form.",
        pd.DataFrame(strobe_rows, columns=["Item", "Reporting point", "Manuscript location", "Status"]),
        "Note: the study is an ecological district-level analysis; STROBE items were mapped to the relevant observational-reporting analogue.",
    )

    sensitivity = pd.read_csv(ROOT / "outputs" / "tables" / "submission_sensitivity_checks.csv")
    sensitivity_display = sensitivity.rename(
        columns={
            "Spatial_CV_AUC_mean": "Spatial AUC mean",
            "Spatial_CV_AUC_SD": "Spatial AUC SD",
            "N_folds": "Folds",
        }
    )
    write_table_docx(
        PKG / "Sensitivity_Checks_Tabulated.docx",
        "Sensitivity Checks - Tabulated Output",
        "Reviewer-facing sensitivity checks added to keep the ecological and spatial assumptions visible.",
        sensitivity_display,
    )

    qa_rows = [
        ["Scientific contribution", "New Ghana district-level HIV-TB spatial-ML synthesis; useful because it quantifies spatial over-optimism in random-fold ML.", "Strong"],
        ["Core spatial result", "TB-HIV co-infection Moran's I = 0.469; 50 LISA High-High districts; 48 bivariate High-High districts.", "Strong"],
        ["Prediction claim", "Use leave-one-region-out spatial AUC 0.798 as the main validation result; random 10-fold AUC 0.998 is retained as a cautionary contrast.", "Defensible"],
        ["GWR result", "Global R2 = 0.917; mean local R2 = 0.854, with spatial non-stationarity stated clearly.", "Strong"],
        ["Sensitivity layer", "Removing DHS behavioural/HIV predictors lowers spatial AUC to 0.604-0.655 but leaves a structural/TB/system signal.", "Moderate"],
        ["Guan convention", "Guan and Krachi East remain Low-Low and not Gi* significant, so the shared-polygon convention does not create the High-High finding.", "Acceptable"],
        ["Main limitation", "Ecological design and regional-era DHS behavioural inputs limit causal and fine-grained behavioural inference.", "Needs upfront wording"],
        ["Journal fit", "Epidemiology & Infection fit is credible because the paper is infectious-disease surveillance, spatial epidemiology and methods-facing.", "Good"],
        ["Acceptance odds", "Honest estimate after current fixes: 60-68%, not >70%.", "Candid"],
        ["Submission tone", "Lead with spatial honesty, not 'near-perfect ML'.", "Ready"],
    ]
    write_table_docx(
        PKG / "Tabulated_QA_Readiness_Output.docx",
        "Tabulated QA and Editorial Readiness Output",
        "Cambridge-style editor-facing QA summary, humanised and aligned with the final dashboard, poster, repository and manuscript package.",
        pd.DataFrame(qa_rows, columns=["Domain", "Finding", "Readiness"]),
        "Bottom line: S+ package quality for coherence and presentation; acceptance remains constrained by ecological data limitations.",
    )

    qa = """# Package QA Note

At package level, this is now S+: the manuscript, HI-EI dashboard, poster, repository metadata, cover letter, STROBE location checklist and data/code statements now tell the same story.

Do not lead with "near-perfect ML". Lead with spatial honesty. HIV-TB co-infection clusters geographically, and the headline model falls from 0.998 under random 10-fold CV to 0.798 under leave-one-region-out CV. That is the publishable argument.

One sensitivity package has been added at `outputs/tables/submission_sensitivity_checks.csv`. Structural/TB/system-only predictors retain weaker but real leave-one-region-out discrimination (best AUC 0.655). Guan and Krachi East are both Low-Low and not Gi* significant, so the shared-polygon convention is not creating the reported High-High finding.

Honest acceptance estimate for Epidemiology & Infection: 60-68%. A claim above 70% would need newer district-level behavioural/HIV inputs, or an editor who is especially receptive to ecological spatial-ML work.

The visual artefacts were regenerated through the Bespoke HI-EI generator. They now carry the corrected numbers: spatial AUC 0.798, 50 LISA High-High districts, 48 bivariate High-High districts, GWR global R2 0.917 / mean local R2 0.854, and the actual 260districts repository slug.
"""
    write_simple_docx(PKG / "Humanised_QA_Readiness_Note.docx", qa)


def main() -> None:
    PKG.mkdir(exist_ok=True)
    text = read_manuscript()
    md_to_docx(text)
    write_cover_letter()
    write_checklists()
    for markdown_file in PKG.glob("*.md"):
        markdown_file.unlink()
    print(f"Wrote {PKG}")
    print(f"Manuscript DOCX: {MANUSCRIPT_DOCX} ({MANUSCRIPT_DOCX.stat().st_size / 1024 / 1024:.1f} MB)")


if __name__ == "__main__":
    main()
